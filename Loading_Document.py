from pathlib import Path
from io import BytesIO
import os
from typing import List

import pandas as pd
from dotenv import load_dotenv

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_openai import OpenAIEmbeddings
from langchain_openai import ChatOpenAI
from langchain_community.vectorstores import FAISS

# load environment variables (OPENAI_API_KEY or API_Agents_Key)
load_dotenv()
API_KEY = os.getenv("API_Agents_Key1") or os.getenv("OPENAI_API_KEY")
if API_KEY:
    os.environ["OPENAI_API_KEY"] = API_KEY

def excel_to_documents(excel_path: Path) -> List[Document]:
    """
    Read all sheets from an Excel file and convert rows into plain-text Documents.
    """
    xlsx = pd.read_excel(excel_path, sheet_name=None, dtype=str)
    docs: List[Document] = []
    for sheet_name, df in xlsx.items():
        df = df.fillna("")
        for i, row in df.iterrows():
            parts = [f"Sheet: {sheet_name}", f"Row: {i}"]
            for col, val in row.items():
                parts.append(f"{col}: {val}")
            text = "\n".join(parts)
            docs.append(Document(page_content=text, metadata={"source": str(excel_path), "sheet": sheet_name}))
    return docs

def textfile_to_documents(text_path: Path) -> List[Document]:
    text = text_path.read_text(encoding="utf8", errors='ignore')
    return [Document(page_content=text, metadata={"source": str(text_path)})]

def pdf_to_documents(pdf_path: Path) -> List[Document]:
    """Load PDF files using PyPDF"""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(str(pdf_path))
        docs = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text.strip():
                docs.append(Document(
                    page_content=text,
                    metadata={"source": str(pdf_path), "page": i + 1}
                ))
        return docs if docs else [Document(page_content="", metadata={"source": str(pdf_path)})]
    except ImportError:
        # Fallback: treat as text
        return textfile_to_documents(pdf_path)

def docx_to_documents(docx_path: Path) -> List[Document]:
    """Load DOCX files using python-docx"""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(docx_path))
        text = "\n".join([para.text for para in doc.paragraphs])
        return [Document(page_content=text, metadata={"source": str(docx_path)})]
    except ImportError:
        # Fallback: treat as text
        return textfile_to_documents(docx_path)

def csv_to_documents(csv_path: Path) -> List[Document]:
    """Load CSV files similar to Excel"""
    df = pd.read_csv(csv_path, dtype=str)
    df = df.fillna("")
    docs = []
    for i, row in df.iterrows():
        parts = [f"Row: {i}"]
        for col, val in row.items():
            parts.append(f"{col}: {val}")
        text = "\n".join(parts)
        docs.append(Document(page_content=text, metadata={"source": str(csv_path)}))
    return docs

def load_source_as_documents(path: str) -> List[Document]:
    """
    Load documents from various file types.
    Supports: Excel (.xlsx, .xls), PDF (.pdf), Word (.docx), CSV (.csv), Text (.txt, .md, etc.)
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"Source not found: {p}")
    
    suffix = p.suffix.lower()
    
    # Route to appropriate loader based on file extension
    if suffix in (".xlsx", ".xls"):
        return excel_to_documents(p)
    elif suffix == ".pdf":
        return pdf_to_documents(p)
    elif suffix == ".docx":
        return docx_to_documents(p)
    elif suffix == ".csv":
        return csv_to_documents(p)
    else:
        # Default: treat as text file
        return textfile_to_documents(p)

def build_faiss_index(
    source_path: str,
    persist_dir: str,
    chunk_size: int = 500,
    chunk_overlap: int = 50,
) -> None:
    """
    Build a FAISS index from a source file (Excel or text) and persist it to disk.
    """
    docs = load_source_as_documents(source_path)

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", " ", ""],
    )
    chunks = splitter.split_documents(docs)

    embeddings = OpenAIEmbeddings(
        openai_api_key=API_KEY,
        openai_api_base="https://openrouter.ai/api/v1",
        model="openai/text-embedding-3-small"
    )
    vectorstore = FAISS.from_documents(chunks, embeddings)

    persist_path = Path(persist_dir)
    persist_path.mkdir(parents=True, exist_ok=True)
    vectorstore.save_local(str(persist_path))

def load_faiss_index(persist_dir: str) -> FAISS:
    """
    Load a persisted FAISS index from disk.
    """
    embeddings = OpenAIEmbeddings(
        openai_api_key=API_KEY,
        openai_api_base="https://openrouter.ai/api/v1",
        model="openai/text-embedding-3-small"
    )
    return FAISS.load_local(persist_dir, embeddings, allow_dangerous_deserialization=True)

def query_faiss(
    query: str,
    persist_dir: str,
    k: int = 4,
    llm_temperature: float = 0.0,
) -> str:
    """
    Run a retrieval query against a persisted FAISS index.
    Returns the LLM answer (string).
    """
    from langchain_core.prompts import ChatPromptTemplate
    from langchain_core.output_parsers import StrOutputParser
    from langchain_core.runnables import RunnablePassthrough
    
    vectorstore = load_faiss_index(persist_dir)
    retriever = vectorstore.as_retriever(search_type="similarity", search_kwargs={"k": k})
    
    # Get relevant documents using invoke (new LangChain API)
    docs = retriever.invoke(query)
    context = "\n\n".join([doc.page_content for doc in docs])
    
    # Create prompt template
    template = """Answer the question based only on the following context:

{context}

Question: {question}

Answer:"""
    
    prompt = ChatPromptTemplate.from_template(template)
    llm = ChatOpenAI(
        model_name="google/gemini-2.0-flash-001",
        openai_api_key=API_KEY,
        openai_api_base="https://openrouter.ai/api/v1",
        temperature=llm_temperature,
        model_kwargs={
            "extra_headers": {
                "HTTP-Referer": "https://localhost",
                "X-Title": "HR Chatbot"
            }
        }
    )
    
    # Create chain
    chain = (
        {"context": lambda x: context, "question": RunnablePassthrough()}
        | prompt
        | llm
        | StrOutputParser()
    )
    
    return chain.invoke(query)

if __name__ == "__main__":
    # quick CLI test
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", required=True, help="Path to .txt or .xlsx source file")
    parser.add_argument("--persist", required=True, help="Directory to persist FAISS index")
    parser.add_argument("--query", help="Run a sample query after building")
    args = parser.parse_args()

    build_faiss_index(args.source, args.persist)
    print("FAISS index built and saved to", args.persist)
    if args.query:
        print("Query result:", query_faiss(args.query, args.persist))